from langchain.document_loaders import PyPDFLoader, PDFMinerLoader, DirectoryLoader
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from os.path import join
import os
from dotenv import load_dotenv
# os.environ['OPENAI_API_KEY']= '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'
load_dotenv(r'C:\Users\sksha\Desktop\llm-assignment-master\llm-assignment-master\llm-assignment-master_\backend\.env')
openai_api_key = os.environ.get('OPENAI_API_KEY')
from langchain.document_loaders import TextLoader, PDFMinerLoader, UnstructuredWordDocumentLoader, CSVLoader

# def load_documents(file_path):
#     if file_path.endswith('.txt'):
#         loader = TextLoader(file_path)
#     elif file_path.endswith('.pdf'):
#         loader = PyPDFLoader(file_path)
#     elif file_path.endswith('.doc') or file_path.endswith('.docx'):
#         loader = UnstructuredWordDocumentLoader(file_path)
#     elif file_path.endswith('.csv'):
#         loader = CSVLoader(file_path)
#     else:
#         raise ValueError(f"Unsupported file format: {file_path}")

#     documents = loader.load()
#     return documents
from fastapi import UploadFile
from typing import List
import fitz  # PyMuPDF
import pandas as pd
import docx
from langchain.docstore.document import Document
def read_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def read_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def read_csv(file_path: str) -> str:
    df = pd.read_csv(file_path)
    return df.to_string()

def read_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

async def load_documents(file: UploadFile)->List[Document]:
    temp_file_path = f"temp_{file.filename}"
    try:
        # Save the uploaded file to a temporary file
        with open(temp_file_path, "wb") as temp_file:
            temp_file.write(await file.read())

        content = ""
        if file.filename.endswith('.pdf'):
            content = read_pdf(temp_file_path)
        elif file.filename.endswith('.docx'):
            content = read_docx(temp_file_path)
        elif file.filename.endswith('.csv'):
            content = read_csv(temp_file_path)
        elif file.filename.endswith('.txt'):
            content = read_txt(temp_file_path)
        else:
            raise ValueError("Unsupported file format")
    except Exception as e:
        # Handle general errors - log or adjust as necessary for your application
        print(f"Error processing document: {e}")
        content = "Error processing document."
    finally:
        # Cleanup: remove the temporary file
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

    metadata = {'source': file.filename}
    document = Document(page_content=content, metadata=metadata)
    return [document]


from langchain.text_splitter import CharacterTextSplitter

def chunk_documents(documents, chunk_size=1000, chunk_overlap=200):
    text_splitter = CharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    chunked_docs = text_splitter.split_documents(documents)
    return chunked_docs


from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma

def create_embeddings(chunked_docs, collection_name):
    embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
    vector_store = Chroma.from_documents(chunked_docs, embeddings, collection_name=collection_name)
    vector_store.persist()

    return vector_store